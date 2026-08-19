# -*- coding: utf-8 -*-
"""
tax_builder.py — 数据驱动生成「完整退税资料」全套
输入：退税资料.xlsx（结构见 export_template 导出的模板）
输出：报关单 / 申报要素(按品名) / PI单 / 采购合同(按供方) / 出库单(按供方)

规则（沿用已确认口径）：
  - 退税类型 R1 决定模式：含「泰国正报」或「正清进泰顺」→ th(美金)；含「香港正报」或「灰清进泰国」→ hk(人民币)
  - 报关单换算：
        hk：单价 = 合同单价/1.13*1.125            （人民币）
        th：单价 = 合同单价/1.13*1.125/当月首工作日美金汇率（转美元，币制「美金」）
    总价 = 换算单价 * 数量；单价保留2位，缩进(floor/ceil)使与合同总价差异最少
        hk 对比人民币合同总价；th 对比美元换算总价（避免量纲错位）
  - 报关单：商品行「数量及单位」下方写净重（kg，来自退税类型表），不写「净重」二字
        表头毛重/净重栏填真实合计；出口日期/申报日期/提运单号/报关人员电话留空
  - 商品行宋体6.5pt(pStyle=5)；单价/总价/币制三行(w:br)分行
  - PI单：不填 Color 列
  - 申报要素：标题30pt bold + 空行 + 序号18.5pt + 内容14pt（宋体）；置物架补0-5序号，毛巾架/地板刷用自带0-4序号
"""
import os
import re
import math
import copy
import shutil
import datetime
import openpyxl
from openpyxl.styles import Font, Alignment
import docx
from docx import shared
from docx.oxml.ns import qn

sys_enc = lambda: None
try:
    import sys
    sys.stdout.reconfigure(encoding='utf-8')
except Exception:
    pass

# ===== 模板路径（全部基于用户指定的输出模版：桌面「提柜时间+柜号+退税类型」，已复制到项目内）=====
HERE = os.path.dirname(os.path.abspath(__file__))
_TPL_DIR = os.path.join(HERE, 'templates', '参考模版')
CUSTOMS_TPL = os.path.join(_TPL_DIR, '提柜时间+-+报关单+-+柜号.docx')
PI_TPL = os.path.join(HERE, 'templates', 'PI_MSMU5432528.xlsx')  # 原始未填 PI 模版（build_pi 按此结构写入，参考模版的已填版含合并单元格不兼容）
DECLARE_TPL = os.path.join(_TPL_DIR, '品名+申报要素.docx')
# 合同+出库单 空白模板（用户提供的「退税资料模版.xlsx」里的「合同和出库单」sheet）
CONTRACT_TPL = os.path.join(HERE, 'templates', '退税资料模版.xlsx')

# ===== 参考输出模版（用户指定：桌面「提柜时间+柜号+退税类型」文件夹，已复制到项目内）=====
# 这些文件本身是「输出模版」：文件名是占位符，内部标红处是示例值，生成时复制并替换标红值。
REF_DIR = os.path.join(HERE, 'templates', '参考模版')
REF_CONTRACT = os.path.join(REF_DIR, '供应商简称+采购合同+合同编号.docx')
REF_OUTBOUND = os.path.join(REF_DIR, '供应商简称+出库单+合同编号.xlsx')

RED_COLORS = {'FFFF0000', 'FF0000', 'red', 'FF0000FF'}


def _strip_red(run):
    """去掉 run 的标红颜色（示例值），恢复默认黑字。"""
    try:
        rPr = run._r.get_or_add_rPr()
        color_el = rPr.find(qn('w:color'))
        if color_el is not None:
            rPr.remove(color_el)
    except Exception:
        pass


def _strip_all_red(doc):
    """把 docx 内所有标红 run 恢复黑字（最终兜底，确保无残留标红）。"""
    targets = []
    for p in doc.paragraphs:
        targets.extend(p.runs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    targets.extend(p.runs)
    for run in targets:
        if run.font.color and run.font.color.rgb is not None and str(run.font.color.rgb) in RED_COLORS:
            _strip_red(run)


def _docx_replace(doc, repl_map):
    """对 docx 全部 run 做子串替换；被改动且原标红的 run 去红。
    repl_map: {示例字符串: 替换值}；值为 '' 表示清空。"""
    targets = []
    for p in doc.paragraphs:
        targets.extend(p.runs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    targets.extend(p.runs)
    for run in targets:
        if not run.text:
            continue
        new = run.text
        changed = False
        for old, nv in repl_map.items():
            if old and old in new:
                new = new.replace(old, str(nv))
                changed = True
        if changed:
            run.text = new
            _strip_red(run)


def _center_table_cells(table):
    """把表格所有单元格的段落设为居中对齐(alignment=1=CENTER)。"""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER


def _unify_font(doc, name='\u4eff\u5b8b'):
    """将文档全部 run 的字体统一为 name（中英文均设置），实现“字体样式统一”。
    只改字体族、不改字号，故标题/正文/表头各自大小保持不变。"""
    targets = []
    for p in doc.paragraphs:
        targets.extend(p.runs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    targets.extend(p.runs)
    for run in targets:
        run.font.name = name
        rPr = run._r.get_or_add_rPr()
        rf = rPr.find(qn('w:rFonts'))
        if rf is None:
            rf = rPr.makeelement(qn('w:rFonts'), {})
            rPr.insert(0, rf)
        rf.set(qn('w:eastAsia'), name)
        rf.set(qn('w:ascii'), name)
        rf.set(qn('w:hAnsi'), name)


def _set_cell_font_black(ws, r1, r2, c1, c2):
    """把范围内标红字体改为黑色。"""
    from openpyxl.styles import Font as _F
    for r in range(r1, r2 + 1):
        for c in range(c1, c2 + 1):
            cell = ws.cell(r, c)
            f = cell.font
            if f is not None and f.color is not None and str(f.color.rgb) in RED_COLORS:
                cell.font = _F(name=f.name, sz=f.sz, bold=f.bold, italic=f.italic,
                               color='FF000000')

# 默认买方（广州狮爪，香港模式；泰国模式可在调用时覆盖）
DEFAULT_BUYER = {
    'name': '广州狮爪网络科技有限公司',
    'address': '广州市天河区宦溪西路20号611、612房',
    'signer': '康小彪',
    'bank': '中国银行股份有限公司广州南国花园支行',
    'account': '637974294993',
}


# ===================== 解析 =====================
def parse_tax_type(path):
    """解析退税类型.xlsx → (meta, suppliers)
    meta: {type_label, mode, sign_date(datetime.date), pickup_date(datetime.date)}
    suppliers: [{name, address, signer, bank, account, contract_no, products:[...]}]
    product: {cn, en, hs, qty, unit, price, amount, origin, gross, net, declare_raw}
    """
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.active

    type_label = (ws['B1'].value or '')
    sign_date = ws['B2'].value
    pickup_date = ws['B3'].value
    if isinstance(sign_date, datetime.datetime):
        sign_date = sign_date.date()
    if isinstance(pickup_date, datetime.datetime):
        pickup_date = pickup_date.date()
    # 柜号（全局）：扫描 A 列标签为「柜号」的行（兼容表头浮动）
    container_no = ''
    for rr in range(1, min(ws.max_row, 20) + 1):
        if str(ws.cell(rr, 1).value or '').strip() == '柜号':
            container_no = str(ws.cell(rr, 2).value or '').strip()
            break

    label = str(type_label)
    if ('泰国正报' in label) or ('正清进泰顺' in label):
        mode = 'th'
    elif ('香港正报' in label) or ('灰清进泰国' in label):
        mode = 'hk'
    else:
        mode = 'hk'  # 默认香港模式

    meta = {'type_label': label, 'mode': mode, 'sign_date': sign_date,
            'pickup_date': pickup_date, 'container_no': container_no}

    suppliers = []
    cur = None

    def _num2str(v):
        if isinstance(v, bool):
            return str(v)
        if isinstance(v, float):
            return '%d' % int(round(v))
        if isinstance(v, int):
            return str(v)
        return str(v or '')

    # 定位表头行（含「供方名称」）以兼容行号浮动；柜号行也可能在表头之前
    header_row = None
    for rr in range(1, min(ws.max_row, 20) + 1):
        if str(ws.cell(rr, 2).value or '').strip() == '供方名称':
            header_row = rr
            break
    if header_row is None:
        header_row = 4
    for r in range(header_row + 1, ws.max_row + 1):
        name = ws.cell(r, 2).value
        if name not in (None, ''):
            cur = {
                'name': str(name).strip(),
                'address': (ws.cell(r, 3).value or ''),
                'signer': (ws.cell(r, 4).value or ''),
                'bank': (ws.cell(r, 5).value or ''),
                'account': _num2str(ws.cell(r, 6).value),
                'short': str(ws.cell(r, 18).value or '').strip(),
                'products': [],
            }
            suppliers.append(cur)
        cn = ws.cell(r, 7).value
        if not cn:
            continue
        # 英文品名：优先读 Q 列(17)，否则留空（PI 英文列将留空并提示）
        en = ws.cell(r, 17).value or ''
        cur['products'].append({
            'cn': str(cn).strip(),
            'en': str(en).strip(),
            'hs': str(ws.cell(r, 12).value or '').strip(),
            'qty': ws.cell(r, 8).value,
            'unit': (ws.cell(r, 9).value or ''),
            'price': ws.cell(r, 10).value,
            'amount': ws.cell(r, 11).value,
            'origin': (ws.cell(r, 13).value or ''),
            'gross': ws.cell(r, 15).value,
            'net': ws.cell(r, 16).value,
            'cartons': ws.cell(r, 19).value,  # S列：箱数（报关单件数用）
            'declare_raw': ws.cell(r, 14).value or '',
        })

    # 合同号：GZSZ + 签约日期(YYYYMMDD) + 供应商序号(01/02…)
    sd = sign_date or datetime.date.today()
    ymd = '%04d%02d%02d' % (sd.year, sd.month, sd.day)
    for i, s in enumerate(suppliers):
        s['contract_no'] = 'GZSZ-%s%02d' % (ymd, i + 1)

    return meta, suppliers


def parse_declare(raw):
    """申报要素文本 → [(序号, 内容), ...]；无序号则补 0..n-1"""
    lines = [l.strip() for l in (raw or '').split('\n') if l and l.strip()]
    if lines and re.match(r'^\d+[、.．]', lines[0]):
        items = []
        for l in lines:
            m = re.match(r'^(\d+)[、.．]\s*(.*)$', l)
            if m:
                items.append((int(m.group(1)), m.group(2).strip()))
            else:
                items.append((len(items), l))
    else:
        items = [(i, l) for i, l in enumerate(lines)]
    return items


# ===================== 换算 =====================
def _convert(contract_unit, qty, mode='hk', rate=None):
    if mode == 'th':
        if rate is None:
            raise ValueError('泰国正报模式需提供当月第一个工作日美金汇率(rate)')
        raw = contract_unit / 1.13 * 1.125 / rate
    else:
        raw = contract_unit / 1.13 * 1.125
    cands = [math.floor(raw * 100) / 100, math.ceil(raw * 100) / 100]
    target = contract_unit * qty if mode == 'hk' else raw * qty
    best = min(cands, key=lambda u: abs(u * qty - target))
    return round(best, 2), round(best * qty, 2)


# ===================== 文件名辅助 =====================
def _fmt_date_filename(d):
    """输出文件名用的日期：YYYY-MM-DD；None/空 → 占位串。"""
    if not d:
        return '未填提柜时间'
    if isinstance(d, datetime.datetime):
        d = d.date()
    if isinstance(d, datetime.date):
        return '%04d-%02d-%02d' % (d.year, d.month, d.day)
    return str(d).strip() or '未填提柜时间'


def _safe(s):
    """去除文件名非法字符并去首尾空格。"""
    return re.sub(r'[\\/:*?"<>|]', '', str(s or '')).strip()


# ===================== 报关单公共填写 =====================
def _fill_head(table, label, contract_nos_str, pkg_count):
    row4 = table.rows[4]
    tc_list = row4._tr.findall(qn('w:tc'))
    vals = [label, contract_nos_str]
    for idx, tc in enumerate(tc_list[:3]):
        # 保留原段落 pPr（含 pStyle），新段落才能继承表格样式（字体大小），避免重建后字体过大
        orig_p = tc.find(qn('w:p'))
        pPr_src = orig_p.find(qn('w:pPr')) if orig_p is not None else None
        for p in tc.findall(qn('w:p')):
            tc.remove(p)
        new_p = tc.makeelement(qn('w:p'), {})
        if pPr_src is not None:
            new_p.append(copy.deepcopy(pPr_src))
        new_r = new_p.makeelement(qn('w:r'), {})
        rPr = new_r.makeelement(qn('w:rPr'), {})
        rFonts = new_r.makeelement(qn('w:rFonts'), {})
        rFonts.set(qn('w:eastAsia'), '宋体')
        rPr.append(rFonts)
        new_r.append(rPr)
        new_t = new_r.makeelement(qn('w:t'), {})
        new_t.text = vals[idx] if idx < len(vals) else ''
        new_t.set(qn('xml:space'), 'preserve')
        new_r.append(new_t)
        new_p.append(new_r)
        tc.append(new_p)
    row5 = table.rows[5]
    tc5 = row5._tr.findall(qn('w:tc'))
    for tc in tc5:
        text = ''.join(t.text or '' for t in tc.iter(qn('w:t')))
        if '件数' in text:
            # 替换 tc 内数字占位（模板中如「1280」）；无数字则追加值
            replaced = False
            for p in tc.findall(qn('w:p')):
                for r_el in p.findall(qn('w:r')):
                    for t_el in r_el.findall(qn('w:t')):
                        t = t_el.text or ''
                        if re.search(r'\d', t):
                            t_el.text = re.sub(r'\d+', str(pkg_count), t)
                            replaced = True
            if not replaced:
                p = tc.find(qn('w:p'))
                if p is not None:
                    new_r = p.makeelement(qn('w:r'), {})
                    rPr = new_r.makeelement(qn('w:rPr'), {})
                    rFonts = new_r.makeelement(qn('w:rFonts'), {})
                    rFonts.set(qn('w:eastAsia'), '宋体')
                    rPr.append(rFonts)
                    new_r.append(rPr)
                    new_t = new_r.makeelement(qn('w:t'), {})
                    new_t.text = str(pkg_count)
                    new_t.set(qn('xml:space'), 'preserve')
                    new_r.append(new_t)
                    p.append(new_r)


def _fill_weights(table, gross_total=None, net_total=None):
    """表头毛重/净重栏：按标签行(件数/毛重/净重…)定位数值行对应列填写（兼容模板占位值变动）。"""
    row5 = table.rows[5]
    tcs5 = row5._tr.findall(qn('w:tc'))
    g_idx = n_idx = None
    for i, tc in enumerate(tcs5):
        text = ''.join(t.text or '' for t in tc.iter(qn('w:t')))
        if '毛重' in text:
            g_idx = i
        elif '净重' in text:
            n_idx = i
    row6 = table.rows[6]
    tcs6 = row6._tr.findall(qn('w:tc'))
    for idx, val in ((g_idx, gross_total), (n_idx, net_total)):
        if idx is None or idx >= len(tcs6):
            continue
        tc = tcs6[idx]
        orig_p = tc.find(qn('w:p'))
        pPr_src = orig_p.find(qn('w:pPr')) if orig_p is not None else None
        for p in tc.findall(qn('w:p')):
            tc.remove(p)
        new_p = tc.makeelement(qn('w:p'), {})
        if pPr_src is not None:
            new_p.append(copy.deepcopy(pPr_src))
        new_r = new_p.makeelement(qn('w:r'), {})
        rPr = new_r.makeelement(qn('w:rPr'), {})
        rFonts = new_r.makeelement(qn('w:rFonts'), {})
        rFonts.set(qn('w:eastAsia'), '宋体')
        rPr.append(rFonts)
        new_r.append(rPr)
        new_t = new_r.makeelement(qn('w:t'), {})
        new_t.text = str(int(val)) if val is not None else ''
        new_t.set(qn('xml:space'), 'preserve')
        new_r.append(new_t)
        new_p.append(new_r)
        tc.append(new_p)


def _clone_row(table, src_idx):
    src = table.rows[src_idx]._tr
    new = copy.deepcopy(src)
    src.addnext(new)
    return table.rows[src_idx + 1]


def _fill_origin_cell(tc, origin):
    """商品行末端的合并单元格（原产国（地区）/最终目的国(地区)/境内货源地/征免）
    由多个 run 拼接而成：第 3 个非空 w:t 即『境内货源地』字段。
    把该字段替换为真实货源地，使货源地随品名一起更新。"""
    texts = []
    for p in tc.findall(qn('w:p')):
        for r in p.findall(qn('w:r')):
            for t_el in r.findall(qn('w:t')):
                if (t_el.text or '').strip():
                    texts.append(t_el)
    if len(texts) >= 3:
        texts[2].text = (origin or '').strip()


def _write_goods_row(table, r_idx, row, rate):
    """row=(项号,HS,品名,数量及单位,合同单价,数量,模式,币制,净重kg,货源地)"""
    no, hs, name, qty_unit, contract_unit, qty, mode, currency, net_weight, origin = row
    try:
        price, amount = _convert(contract_unit, qty, mode=mode, rate=rate)
    except ValueError:
        price = amount = None
    row_el = table.rows[r_idx]
    tcs = row_el._tr.findall(qn('w:tc'))
    cell0 = f'{no}   {hs}  {name}' if hs else f'{no}   {name}'
    cell_content = [cell0, qty_unit, None]
    for idx, tc in enumerate(tcs[:3]):
        for p in tc.findall(qn('w:p')):
            tc.remove(p)
        new_p = tc.makeelement(qn('w:p'), {})
        pPr = new_p.makeelement(qn('w:pPr'), {})
        pStyle = new_p.makeelement(qn('w:pStyle'), {})
        pStyle.set(qn('w:val'), '5')
        pPr.append(pStyle)
        new_p.append(pPr)
        if idx == 2:
            if price is not None:
                for j, v in enumerate([price, amount, currency]):
                    new_r = new_p.makeelement(qn('w:r'), {})
                    rPr = new_r.makeelement(qn('w:rPr'), {})
                    rFonts = new_r.makeelement(qn('w:rFonts'), {})
                    rFonts.set(qn('w:hint'), 'eastAsia')
                    rPr.append(rFonts)
                    new_r.append(rPr)
                    new_t = new_r.makeelement(qn('w:t'), {})
                    new_t.text = ('%.2f' % v) if j < 2 else v
                    new_t.set(qn('xml:space'), 'preserve')
                    new_r.append(new_t)
                    new_p.append(new_r)
                    if j < 2:
                        br = new_p.makeelement(qn('w:br'), {})
                        br.set(qn('w:type'), 'textWrapping')
                        br_r = new_p.makeelement(qn('w:r'), {})
                        br_r.append(br)
                        new_p.append(br_r)
        else:
            new_r = new_p.makeelement(qn('w:r'), {})
            rPr = new_r.makeelement(qn('w:rPr'), {})
            rFonts = new_r.makeelement(qn('w:rFonts'), {})
            rFonts.set(qn('w:hint'), 'eastAsia')
            rPr.append(rFonts)
            new_r.append(rPr)
            new_t = new_r.makeelement(qn('w:t'), {})
            new_t.text = cell_content[idx]
            new_t.set(qn('xml:space'), 'preserve')
            new_r.append(new_t)
            new_p.append(new_r)
            if idx == 1 and net_weight is not None:
                br = new_p.makeelement(qn('w:br'), {})
                br.set(qn('w:type'), 'textWrapping')
                br_r = new_p.makeelement(qn('w:r'), {})
                br_r.append(br)
                new_p.append(br_r)
                n_r = new_p.makeelement(qn('w:r'), {})
                n_rPr = n_r.makeelement(qn('w:rPr'), {})
                n_rFonts = n_r.makeelement(qn('w:rFonts'), {})
                n_rFonts.set(qn('w:hint'), 'eastAsia')
                n_rPr.append(n_rFonts)
                n_r.append(n_rPr)
                n_t = n_r.makeelement(qn('w:t'), {})
                n_t.text = '%d 千克' % int(net_weight)
                n_t.set(qn('xml:space'), 'preserve')
                n_r.append(n_t)
                new_p.append(n_r)
        tc.append(new_p)
    # 境内货源地：更新末端合并单元格（原产国/最终目的国/境内货源地/征免）中的货源地字段
    if tcs:
        last = tcs[-1]
        if '照章征税' in ''.join(t.text or '' for t in last.iter(qn('w:t'))):
            _fill_origin_cell(last, origin)


# ===================== 生成：报关单 =====================
def build_customs(meta, suppliers, out_dir, th_rate=None, out_name=None):
    if out_name is None:
        # 文件名：提柜时间(复用提货日期)-报关单-柜号.docx
        out_name = '%s-报关单-%s.docx' % (_fmt_date_filename(meta.get('pickup_date')),
                                          _safe(meta.get('container_no')) or '未填柜号')
    out_path = os.path.join(out_dir, out_name)
    shutil.copy(CUSTOMS_TPL, out_path)
    doc = docx.Document(out_path)
    table = doc.tables[0]
    mode = meta['mode']
    currency = '美金' if mode == 'th' else '人民币'
    contract_nos = [s['contract_no'] for s in suppliers]
    all_prod = [(s, p) for s in suppliers for p in s['products']]
    # 件数（总箱数）：优先取表格「箱数」列总和；整列未填则退回商品数量总和
    carton_vals = [float(p.get('cartons') or 0) for s, p in all_prod]
    if any(c > 0 for c in carton_vals):
        pkg_count = int(sum(carton_vals))
    else:
        pkg_count = sum(int(p['qty']) for s, p in all_prod)
    _fill_head(table, '合同协议号', '、'.join(contract_nos), pkg_count)
    gross_total = sum(float(p['gross'] or 0) for s, p in all_prod)
    net_total = sum(float(p['net'] or 0) for s, p in all_prod)
    _fill_weights(table, gross_total, net_total)
    if mode == 'th' and th_rate is None:
        print('提示：泰国正报模式未提供美金汇率，报关单单价/总价列留空，提供汇率后重跑即可补入。')
    n = len(all_prod)
    # 商品行区：表头行（含「项号」「商品编号」）下一行起，至表尾行（含「报关人员/申报单位」）止。
    # 模板通常只有 1 行完整商品示例行，其后为空行/表尾行；统一只保留第一个商品行并按需克隆，
    # 保证每行都具备完整列结构（品名/海关编码/数量/单价/币制/原产国），且不覆盖表尾行。
    hdr_i = None
    for ri, row in enumerate(table.rows):
        txt = ''.join(t.text or '' for t in row._tr.iter(qn('w:t')))
        if '项号' in txt and '商品编号' in txt:
            hdr_i = ri
            break
    hdr_i = 9 if hdr_i is None else hdr_i
    tail_i = len(table.rows)
    for ri in range(hdr_i + 1, len(table.rows)):
        txt = ''.join(t.text or '' for t in table.rows[ri]._tr.iter(qn('w:t')))
        if '报关人员' in txt or '申报单位' in txt:
            tail_i = ri
            break
    first = hdr_i + 1
    # 删除商品区内除第一行外的多余行（从后往前删，索引稳定）
    for ri in range(tail_i - 1, first, -1):
        tr = table.rows[ri]._tr
        tr.getparent().remove(tr)
    # 克隆第一个商品行补足到 n 行
    cur = first
    for _ in range(max(n - 1, 0)):
        _clone_row(table, cur)
        cur += 1
    for i, (s, p) in enumerate(all_prod):
        row = (i + 1, p['hs'], p['cn'], '%s%s' % (int(p['qty']), p['unit']),
               float(p['price']), int(p['qty']), mode, currency, float(p['net'] or 0),
               p.get('origin') or '')
        _write_goods_row(table, first + i, row, th_rate)
    doc.save(out_path)
    return out_path


# ===================== 生成：申报要素 =====================
def build_declare(suppliers, out_dir):
    doc = docx.Document(DECLARE_TPL)
    seen = []
    results = []
    for s in suppliers:
        for p in s['products']:
            if p['cn'] in seen:
                continue
            seen.append(p['cn'])
            items = parse_declare(p['declare_raw'])
            out_path = os.path.join(out_dir, '%s+申报要素.docx' % p['cn'])
            d = docx.Document(DECLARE_TPL)
            body = d.element.body
            for para in list(body.findall(qn('w:p'))):
                body.remove(para)

            def _run(par, text, size, bold=False):
                r = par.add_run(text)
                r.font.name = '宋体'
                r.font.size = shared.Pt(size)
                r.font.bold = bold
                rPr = r._r.get_or_add_rPr()
                rf = rPr.find(qn('w:rFonts'))
                if rf is None:
                    rf = rPr.makeelement(qn('w:rFonts'), {})
                    rPr.insert(0, rf)
                rf.set(qn('w:eastAsia'), '宋体')
                return r

            p0 = d.add_paragraph()
            p0.alignment = 1  # CENTER 文本标题居中
            _run(p0, '申报要素', 30, bold=True)
            d.add_paragraph()
            for idx, content in items:
                par = d.add_paragraph()
                _run(par, '%d、' % idx, 18.5)
                _run(par, content, 14)
            d.save(out_path)
            results.append(out_path)
    return results


# ===================== 生成：PI 单 =====================
def build_pi(meta, suppliers, out_dir, th_rate=None, out_name=None):
    if out_name is None:
        # 文件名：提柜时间(复用提货日期)-PI单-柜号.xlsx
        out_name = '%s-PI单-%s.xlsx' % (_fmt_date_filename(meta.get('pickup_date')),
                                        _safe(meta.get('container_no')) or '未填柜号')
    out_path = os.path.join(out_dir, out_name)
    shutil.copy(PI_TPL, out_path)
    wb = openpyxl.load_workbook(out_path)
    ws = wb['sales contract']
    contract_nos = [s['contract_no'] for s in suppliers]
    ws['J11'] = 'Contract No: ' + ', '.join(contract_nos)

    mode = meta['mode']
    all_prod = [(s, p) for s in suppliers for p in s['products']]
    n = len(all_prod)
    first = 13
    if n > 4:
        ws.insert_rows(17, n - 4)  # 在 R16 与 R18 之间插入，总额公式行随之下移
    for i, (s, p) in enumerate(all_prod):
        r = first + i
        qty = int(p['qty'])
        # PI 单价/总价须与报关单一致：采用同一换算公式（合同单价/1.13*1.125，泰国再除汇率）
        try:
            dc_unit, dc_total = _convert(float(p['price']), qty, mode=mode, rate=th_rate)
        except ValueError:
            dc_unit, dc_total = float(p['price']), float(p['price']) * qty
        ws.cell(row=r, column=1, value=p['hs'])
        ws.cell(row=r, column=4, value=p['en'])  # D: 英文品名（无则留空）
        # column 7 (G) Color 不填
        ws.cell(row=r, column=8, value=qty)                              # H: Qty
        ws.cell(row=r, column=9, value=round(dc_unit, 2))                # I: Price（报关单换算单价）
        ws.cell(row=r, column=15, value=round(dc_total, 2))              # O: Amount（换算总价，与报关单一致）
        ws.cell(row=r, column=20, value='%s(%s)' % (p['cn'], p['unit']))  # 备注

    # 清除模板残留的示例商品行（仅保留 n 个真实商品行），避免 PI 出现非本批品名
    total_row_idx = None
    for r in range(1, ws.max_row + 1):
        if str(ws.cell(r, 1).value or '').strip().upper().startswith('TOTAL'):
            total_row_idx = r
            break
    if total_row_idx is None:
        total_row_idx = (18 if n <= 4 else 14 + n)
    for r in range(first + n, total_row_idx):
        for c in (1, 4, 8, 9, 15, 20):
            ws.cell(row=r, column=c, value='')

    last = first + n - 1
    if n <= 4:
        total_row, bal_row = 18, 19
    else:
        total_row, bal_row = 14 + n, 15 + n
    ws.cell(row=total_row, column=8, value='=SUM(H%d:H%d)' % (first, last))
    ws.cell(row=total_row, column=15, value='=SUM(O%d:O%d)' % (first, last))
    ws.cell(row=bal_row, column=5, value='=O%d' % total_row)
    wb.save(out_path)
    return out_path


# ===================== 生成：采购合同（按供方，从零生成） =====================
def _set_cell(cell, text, name='宋体', size=10.5, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.name = name
    r.font.size = shared.Pt(size)
    r.font.bold = bold
    rPr = r._r.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rf)
    rf.set(qn('w:eastAsia'), name)


def _fmt_date(d):
    if not d:
        return ''
    if isinstance(d, datetime.datetime):
        d = d.date()
    return '%d年%d月%d日' % (d.year, d.month, d.day)


def build_contracts(meta, suppliers, out_dir, buyer=None):
    """复制「供应商简称+采购合同+合同编号.docx」底版，把标红示例值替换为导入数据。
    文件名：供应商简称+采购合同+合同编号.docx。"""
    buyer = buyer or DEFAULT_BUYER
    sd = _fmt_date(meta['sign_date'])
    pd = _fmt_date(meta.get('pickup_date'))
    results = []
    for s in suppliers:
        products = s['products']
        total = round(sum(float(p['amount'] or 0) for p in products), 2)
        short = _safe(s.get('short') or '') or _safe(s['name'])[:4] or s['name']
        out_path = os.path.join(out_dir, '%s+采购合同+%s.docx' % (_safe(short), s['contract_no']))
        shutil.copy(REF_CONTRACT, out_path)
        doc = docx.Document(out_path)
        name = s['name']; cno = s['contract_no']
        addr = (s['address'] or ''); signer = (s['signer'] or '')
        bank = (s['bank'] or ''); account = (s['account'] or '')
        # 示例值 -> 导入数据（底版即桌面参考格式，示例值落在独立 run 内，子串替换可靠）
        repl = {
            '\u798f\u5efa\u53cb\u8c0a\u80f6\u7c98\u5e26\u96c6\u56e2\u6709\u9650\u516c\u53f8': name,
            '\u798f\u5efa\u7701\u798f\u5dde\u5e02\u798f\u6e05\u5e02\u6c5f\u9634\u6e2f\u57ce\u7ecf\u6d4e\u533a\u4e1c\u90e8\u7247\u533a': addr,
            '\u6797\u514b\u5174': signer,
            '\u4e2d\u56fd\u94f6\u884c\u80a1\u4efd\u6709\u9650\u516c\u53f8\u798f\u6e05\u5206\u884c': bank,
            '426058378409': account,
            'GZSZ-2026081401': cno,
        }
        if sd:
            repl['2026\u5e748\u670814\u65e5'] = sd        # 签订日期（供应商信息表“日期：”）
        if pd:
            repl['2026\u5e748\u670822\u65e5'] = pd        # 交(提)货日期
        if buyer is not DEFAULT_BUYER:
            repl.update({
                '\u5e7f\u5dde\u72ee\u722a\u7f51\u7edc\u79d1\u6280\u6709\u9650\u516c\u53f8': buyer['name'],
                '\u5e7f\u5dde\u5e02\u5929\u6cb3\u533a\u5b97\u5b98\u6eaa\u897f\u8def20\u53f7611\u3001612\u623f': buyer['address'],
                '\u5eb7\u5c0f\u5f6c': buyer['signer'],
                '\u4e2d\u56fd\u94f6\u884c\u80a1\u4efd\u6709\u9650\u516c\u53f8\u5e7f\u5dde\u5357\u56fd\u82b1\u56ed\u652f\u884c': buyer['bank'],
                '637974294993': buyer['account'],
            })
        _docx_replace(doc, repl)
        _strip_all_red(doc)   # 兜底：彻底清除残留标红
        # 商品表：保留表头(r0)、合计(r2)、人民币(r3)；仅把示例产品行(r1)替换为导入产品。
        # 注意：list(table._tbl)[1:] 会把 tblGrid 一并删掉导致表格损坏，故只删示例产品行。
        table = doc.tables[0]
        trs = table._tbl.findall(qn('w:tr'))
        # trs[0]=表头  trs[1]=示例产品  trs[2]=合计  trs[3]=人民币
        prod_tpl_tr = copy.deepcopy(trs[1])
        table._tbl.remove(trs[1])          # 删除示例产品行（合计/人民币行保留，位置自动后移）
        prev = trs[0]
        for _ in range(len(products)):
            new_tr = copy.deepcopy(prod_tpl_tr)
            prev.addnext(new_tr)
            prev = new_tr
        # 此时 rows = [表头, 产品1..N, 合计, 人民币]
        rows = table.rows
        n = len(products)
        for i in range(n):
            cells = rows[1 + i].cells
            p = products[i]
            cells[0].text = str(i + 1)
            cells[1].text = p['cn']
            cells[2].text = str(int(p['qty']))
            cells[3].text = '\u4eba\u6c11\u5e01'
            cells[4].text = '%.2f' % float(p['price'])
            cells[5].text = '%.2f' % float(p['amount'] or 0)
        total_row = rows[1 + n]
        rmb_row = rows[2 + n]
        # 合计/人民币行：仅「总金额」列(索引5)填值，其余列留空（与参考模版结构一致）
        total_row.cells[5].text = '%.2f' % total
        total_row.cells[0].text = '\u5408\u8ba1'
        rmb_row.cells[5].text = '%.2f' % total
        rmb_row.cells[0].text = '\u4eba\u6c11\u5e01'
        _center_table_cells(table)           # 产品表文字全部居中
        _unify_font(doc, '\u4eff\u5b8b')     # 全文字体统一为仿宋（字号不变）
        _strip_all_red(doc)
        doc.save(out_path)
        results.append(out_path)
    return results


def build_outbounds(meta, suppliers, out_dir):
    """复制「供应商简称+出库单+合同编号.xlsx」底版，把标红示例值替换为导入数据，
    并保留模板原有格式（标题/表头合并、合计/盖章合并、列宽、边框）。
    仅把示例商品行(R4)替换为导入商品（多商品时在 R5 起插入并复制其样式），
    合计/盖章行原样保留。文件名：供应商简称+出库单+合同编号.xlsx。"""
    pickup = meta['pickup_date']
    if isinstance(pickup, datetime.datetime):
        pickup = pickup.date()
    outbound_date = (pickup - datetime.timedelta(days=1)) if pickup else None
    od_str = '%04d-%02d-%02d' % (outbound_date.year, outbound_date.month, outbound_date.day) if outbound_date else ''
    results = []
    for s in suppliers:
        products = s['products']
        n = len(products)
        short = _safe(s.get('short') or '') or _safe(s['name'])[:4] or s['name']
        out_path = os.path.join(out_dir, '%s+出库单+%s.xlsx' % (_safe(short), s['contract_no']))
        shutil.copy(REF_OUTBOUND, out_path)
        wb = openpyxl.load_workbook(out_path)
        ws = wb.active
        name = s['name']
        ws['A1'] = '%s（ 出库单）' % name
        ws['A2'] = '编制单位：　　日期：%s' % od_str   # 出库单不显示编号（文件以合同号命名）
        # 示例商品行在 R4，取其单元格样式用于多商品插入行
        def _style(row_idx):
            d = {}
            for c in range(1, 10):
                cell = ws.cell(row_idx, c)
                d[c] = (cell.font.copy(), cell.fill.copy(), cell.alignment.copy(),
                        cell.border.copy(), cell.number_format)
            return d
        base = _style(4)

        def _fill(rr, p):
            ws.cell(rr, 1, rr - 3)                      # 序号
            ws.cell(rr, 2, p['cn'])                     # 产品名称
            ws.cell(rr, 3, '')                          # 规格 留空
            ws.cell(rr, 4, p['unit'])                   # 单位
            ws.cell(rr, 5, int(p['qty']))               # 数量
            ws.cell(rr, 6, round(float(p['price']), 2)) # 单价
            ws.cell(rr, 7, '=F%d*E%d' % (rr, rr))       # 金额（公式=单价×数量）
        # 第 1 个商品写入 R4（保留模板格式）
        _fill(4, products[0])
        # 多商品：在 R5 起插入并复制 R4 样式
        if n > 1:
            ws.insert_rows(5, n - 1)
            for i in range(1, n):
                rr = 5 + i - 1
                for c in range(1, 10):
                    f, fi, al, bd, nf = base[c]
                    cell = ws.cell(rr, c)
                    cell.font = f; cell.fill = fi; cell.alignment = al
                    cell.border = bd; cell.number_format = nf
                _fill(rr, products[i])
        # 定位合计/盖章行（模板原在 R10/R11，多商品后下移）
        hr = sr = None
        for r in range(4, ws.max_row + 1):
            v = ws.cell(r, 2).value
            if v == '合计':
                hr = r
            if isinstance(v, str) and '盖章' in v:
                sr = r
        last_prod = 4 + n - 1
        if hr:
            ws.cell(hr, 1, '')
            ws.cell(hr, 2, '合计')
            ws.cell(hr, 5, '=SUM(E4:E%d)' % last_prod)
            ws.cell(hr, 7, '=SUM(G4:G%d)' % last_prod)
        if sr:
            ws.cell(sr, 2, '出库单位盖章：')
            ws.cell(sr, 6, '收货单位盖章：')
        # 重新建立合并（openpyxl insert_rows 不会平移合并区域）
        for m in list(ws.merged_cells.ranges):
            ws.unmerge_cells(str(m))
        ws.merge_cells('A1:H1')
        ws.merge_cells('A2:H2')
        if sr:
            ws.merge_cells('I3:I%d' % sr)
            ws.merge_cells('B%d:E%d' % (sr, sr))
            ws.merge_cells('F%d:H%d' % (sr, sr))
        _set_cell_font_black(ws, 4, sr if sr else (hr or 4), 1, 8)
        wb.save(out_path)
        results.append(out_path)
    return results


# ===================== 生成：采购合同 + 出库单（套用空白模板「合同和出库单」） =====================
def fill_contract_outbound(supplier, meta, out_path):
    """套用「退税资料模版.xlsx」的「合同和出库单」sheet，按供应商填值输出。"""
    wb = openpyxl.load_workbook(CONTRACT_TPL)
    ws = wb['合同和出库单']
    products = supplier['products']
    n = len(products)
    sign_date = meta['sign_date']
    pickup_date = meta['pickup_date']
    sname = supplier['name']; saddr = supplier['address']; ssign = supplier['signer']
    sbank = supplier['bank']; sacc = supplier['account']; cno = supplier['contract_no']
    shift = n - 1  # 合同插行导致的整体下移
    # 备份原合并，稍后恢复（避免写冲突先全部解除）
    orig = [(r.min_row, r.min_col, r.max_row, r.max_col) for r in ws.merged_cells.ranges]
    for rng in list(ws.merged_cells.ranges):
        ws.unmerge_cells(str(rng))
    # —— 合同商品区 ——
    cr = None
    for r in range(14, 26):
        if ws.cell(r, 1).value == '合计':
            cr = r
            break
    if n > 1:
        ws.insert_rows(15, n - 1)
        cr = cr + (n - 1)
    for i, p in enumerate(products):
        r = 14 + i
        ws.cell(r, 1, i + 1)
        ws.cell(r, 2, p['cn'])
        ws.cell(r, 3, int(p['qty']))
        ws.cell(r, 5, '人民币')
        ws.cell(r, 6, float(p['price']))
        ws.cell(r, 8, float(p['amount']))
    total_amt = round(sum(float(p['amount'] or 0) for p in products), 2)
    ws.cell(cr, 1, '合计')
    ws.cell(cr, 2, total_amt)
    # —— 供应商身份（替换模板示例）——
    ws.cell(6, 2, sname)
    ws.cell(36, 2, sname)
    ws.cell(37, 2, saddr)
    ws.cell(39, 4, ssign)
    ws.cell(40, 2, sbank)
    ws.cell(42, 2, str(sacc))          # 账号以字符串写入，避免 1e16 科学计数法
    ws.cell(48 + shift, 1, sname)      # 出库单标题（随位移下移）
    # —— 合同号 / 日期 ——
    ws.cell(6, 8, cno)
    ws.cell(7, 8, '广州')
    ws.cell(21, 5, pickup_date)
    ws.cell(43, 2, sign_date)
    ws.cell(43, 7, sign_date)
    # —— 出库单商品区 ——
    oh = None
    for r in range(45, 60):
        if ws.cell(r, 1).value == '序号':
            oh = r
            break
    ocr = None
    for r in range(oh, 60):
        if ws.cell(r, 2).value == '合计':
            ocr = r
            break
    if n > 1:
        ws.insert_rows(oh + 2, n - 1)
        ocr = ocr + (n - 1)
    for i, p in enumerate(products):
        r = oh + 1 + i
        ws.cell(r, 1, i + 1)
        ws.cell(r, 2, p['cn'])
        ws.cell(r, 4, p['unit'])
        ws.cell(r, 5, int(p['qty']))
        ws.cell(r, 6, float(p['price']))
        ws.cell(r, 7, float(p['amount']))
    ws.cell(ocr, 2, '合计')
    ws.cell(ocr, 5, sum(int(p['qty']) for p in products))
    ws.cell(ocr, 7, total_amt)
    # 出库日期 = 提货前1天；模板未填提货日期时留空，避免报错
    if pickup_date:
        ws.cell(49 + shift, 6, pickup_date - datetime.timedelta(days=1))
    else:
        ws.cell(49 + shift, 6, None)
    # —— 恢复模板原合并（跳过重算区）——
    for (r1, c1, r2, c2) in orig:
        if r1 in (14, 15, 16):
            continue
        if c1 == 9 and r1 >= 48:
            continue
        ws.merge_cells(start_row=r1, start_column=c1, end_row=r2, end_column=c2)
    # —— 重算：商品行 / 合计 / 出库单列 合并 ——
    for i in range(n):
        r = 14 + i
        ws.merge_cells(f"C{r}:D{r}")
        ws.merge_cells(f"F{r}:G{r}")
        ws.merge_cells(f"H{r}:I{r}")
    ws.merge_cells(f"B{cr}:I{cr + 1}")
    ws.merge_cells(f"I{oh}:I{ocr + 1}")
    wb.save(out_path)


def build_contract_outbound_all(meta, suppliers, out_dir):
    results = []
    for s in suppliers:
        # 文件名取供方名前 4 字，避免过长
        short = re.sub(r'[\\/:*?"<>|]', '', s['name'])[:4]
        out_path = os.path.join(out_dir, '%s-%s-采购合同和出库单.xlsx' % (s['contract_no'], short))
        fill_contract_outbound(s, meta, out_path)
        results.append(out_path)
    return results


# ===================== 导出「退税资料」模板（纯数据，不套用渲染模板的 Excel 格式） =====================
# 模板结构对齐 parse_tax_type 读取的列：
#   R1  A=退税类型  B=退税模式说明（含「香港正报/灰清进泰国」→ 香港；含「泰国正报/正清进泰顺」→ 泰国；网页选择可覆盖）
#   R2  A=签约日期  B=日期
#   R3  A=提货日期  B=日期（同时用作输出文件名里的「提柜时间」）
#   R4  A=柜号      B=柜号（全局，用于报关单/PI 文件名）
#   R5  表头：合同|供方名称|单位地址|法定代表人或代理人签名|开户行|银行账号|
#              产品名称|数量|单位|单价|总金额|报关编码|货源地|申报要素|毛重|净重|英文品名|供应商简称|箱数
#   R6+ 供方/商品行（B列有值=新供方；G列有值=该供方下的商品）
#   箱数（S列）：每个商品的外箱数，报关单表头「件数」= 各行箱数总和；整列留空则退回商品数量总和
_TEMPLATE_HEADERS = ['合同', '供方名称', '单位地址', '法定代表人或代理人签名', '开户行', '银行账号',
                     '产品名称', '数量', '单位', '单价', '总金额', '报关编码', '货源地',
                     '申报要素', '毛重', '净重', '英文品名', '供应商简称', '箱数']

# 内置示例（与已验证数据一致，便于直接看到填写格式并可一键复算）
# product = (产品名称, 数量, 单位, 单价, 总金额, 报关编码, 货源地, 申报要素, 毛重, 净重, 英文品名, 箱数)
_TEMPLATE_SAMPLE = [
    (1, '台州市好一家家居用品有限公司', '浙江省台州市椒江区洪家街道灵香店园', '廖冬玲',
     '中国农业银行台州支行', '19900001040023700',
     [('置物架', 755, '个', 27.05, 20422.75, '3924900000', '台州',
       '无品牌\n不享惠\n置物用\nPS+PP\n无品牌\n两层、三层', 1500, 1300, '', 63)]),
    (2, '台州黄岩亿杰日用品有限公司', '浙江省台州市黄岩区南城街道药山村1区97号', '胡海权',
     '中国农业银行股份有限公司台州南城支行', '19915301040002400',
     [('毛巾架', 228, '箱', 243, 55404, '3924900000', '台州',
       '0、无品牌\n\n1、不享惠\n\n2、置放物品\n\n3、PP \n\n4、白色', 800, 700, '', 228),
      ('地板刷', 48, '箱', 486, 23328, '9603909090', '',
       '0、无品牌\n\n1、不享惠\n\n2、清洁地板\n\n3、PP 不锈钢、TPR\n\n4、3合1绿黄色、3合1白灰色、3合1深蓝色、4合1白灰色、4合1橙灰色、2合1白色', 650, 500, '', 48),
      ('地板刷', 21, '箱', 402, 8442, '9603909090', '',
       '0、无品牌\n\n1、不享惠\n\n2、清洁地板\n\n3、PP 不锈钢、TPR\n\n4、4合1白灰色、4合1橙灰色', 450, 350, '', 21)]),
]


def export_template(out_path, sample=False, mode='hk'):
    """导出「退税资料」数据模板（纯数据 xlsx，不套用渲染模板的 Excel 合并/样式）。
    sample=True ：带内置示例数据（可直接上传复算）；False（默认）：仅表头 + 占位空行，无示例数据。
    返回 out_path。
    """
    if mode == 'th':
        b1 = '广州出泰国正报，正清进泰顺公司'
    else:
        b1 = '广州出香港正报退税，灰清进泰国（不进公司）'

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = '退税资料'

    ws['A1'] = '退税类型'
    ws['B1'] = b1
    ws['A2'] = '签约日期'
    ws['A3'] = '提货日期'
    ws['A4'] = '柜号'
    if not sample:
        ws['B4'] = ''  # 柜号留空，用户填写（全局，用于报关单/PI 文件名）
    for c, h in enumerate(_TEMPLATE_HEADERS, 1):
        ws.cell(row=5, column=c, value=h)

    if sample:
        ws['B2'] = datetime.date(2026, 5, 14)
        ws['B3'] = datetime.date(2026, 5, 20)
        ws['B4'] = 'MSKU0000000'
        r = 6
        for contract_no, name, addr, signer, bank, acc, prods in _TEMPLATE_SAMPLE:
            first = True
            for (cn, qty, unit, price, amount, hs, origin, declare, gross, net, en, cartons) in prods:
                if first:
                    ws.cell(row=r, column=1, value=contract_no)
                    ws.cell(row=r, column=2, value=name)
                    ws.cell(row=r, column=3, value=addr)
                    ws.cell(row=r, column=4, value=signer)
                    ws.cell(row=r, column=5, value=bank)
                ac = ws.cell(row=r, column=6, value=acc)
                ac.number_format = '@'  # 银行账号按文本，避免科学计数法
                ws.cell(row=r, column=18, value=name[:4])  # 供应商简称（示例取前4字，实际由用户填写）
                first = False
                ws.cell(row=r, column=7, value=cn)
                ws.cell(row=r, column=8, value=qty)
                ws.cell(row=r, column=9, value=unit)
                ws.cell(row=r, column=10, value=price)
                ws.cell(row=r, column=11, value=amount)
                ws.cell(row=r, column=12, value=hs)
                ws.cell(row=r, column=13, value=origin)
                ws.cell(row=r, column=14, value=declare)
                ws.cell(row=r, column=15, value=gross)
                ws.cell(row=r, column=16, value=net)
                ws.cell(row=r, column=17, value=en)
                ws.cell(row=r, column=19, value=cartons)  # S列：箱数（报关单件数）
                r += 1
    else:
        # 完全空白占位（无示例数据）：仅表头 + 占位空行，供用户填写后导入。
        # 规则：B列有值=新供应商；G列有值=该供应商下的商品。
        ws.cell(row=6, column=2, value='（供应商①名称）')
        ac = ws.cell(row=6, column=6, value='（银行账号）'); ac.number_format = '@'
        ws.cell(row=6, column=18, value='（简称①）')
        ws.cell(row=6, column=7, value='（产品①名称）')
        ws.cell(row=6, column=8, value=0); ws.cell(row=6, column=9, value='个')
        ws.cell(row=6, column=10, value=0); ws.cell(row=6, column=11, value=0)
        ws.cell(row=6, column=15, value=0); ws.cell(row=6, column=16, value=0)
        ws.cell(row=6, column=19, value=0)  # 箱数（报关单件数，可留空则按数量合计）
        ws.cell(row=7, column=7, value='（产品②名称）')
        ws.cell(row=7, column=8, value=0); ws.cell(row=7, column=9, value='个')
        ws.cell(row=7, column=10, value=0); ws.cell(row=7, column=11, value=0)
        ws.cell(row=7, column=15, value=0); ws.cell(row=7, column=16, value=0)
        ws.cell(row=7, column=19, value=0)  # 箱数
        # 第二个供应商占位（演示 B列触发新供应商）
        ws.cell(row=8, column=2, value='（供应商②名称）')
        ac2 = ws.cell(row=8, column=6, value='（银行账号）'); ac2.number_format = '@'
        ws.cell(row=8, column=18, value='（简称②）')
        ws.cell(row=8, column=7, value='（产品①名称）')
        ws.cell(row=8, column=8, value=0); ws.cell(row=8, column=9, value='个')
        ws.cell(row=8, column=10, value=0); ws.cell(row=8, column=11, value=0)
        ws.cell(row=8, column=15, value=0); ws.cell(row=8, column=16, value=0)
        ws.cell(row=8, column=19, value=0)  # 箱数

    # 仅做表头加粗 + 居中 + 冻结前 5 行，不做任何合并等 Excel 渲染模板样式
    for c in range(1, len(_TEMPLATE_HEADERS) + 1):
        cell = ws.cell(row=5, column=c)
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    ws.freeze_panes = 'A6'
    wb.save(out_path)
    return out_path


# ===================== 总入口 =====================
def collect_summary(meta, suppliers, mode='hk', rate=None):
    """生成后汇总关键数据，供前端「确认页」展示。返回可 JSON 序列化的 dict。

    mode 以调用方传入为准（网页模式开关可能覆盖退税类型第1行自动判定）。
    """
    cur = mode if mode in ('hk', 'th') else (meta.get('mode') or 'hk')

    def _d(v):
        if v is None or v == '':
            return ''
        if isinstance(v, datetime.datetime):
            return v.strftime('%Y-%m-%d')
        if isinstance(v, datetime.date):
            return v.strftime('%Y-%m-%d')
        if isinstance(v, float):
            return round(v, 2)
        return v

    items = []
    total_amount = 0.0   # 合同金额合计（人民币）
    total_declare = 0.0  # 报关金额合计（换算后：hk 人民币 / th 美金）
    for s in suppliers:
        for p in s['products']:
            qty = p.get('qty') or 0
            price = p.get('price') or 0
            amount = p.get('amount')
            if amount in (None, ''):
                amount = price * qty
            dc_unit, dc_total = _convert(price, qty, cur, rate)
            try:
                total_amount += float(amount)
            except Exception:
                pass
            total_declare += dc_total
            items.append({
                'supplier': s.get('name', ''),
                'short': s.get('short', ''),
                'contract_no': s.get('contract_no', ''),
                'cn': p.get('cn', ''),
                'qty': qty,
                'unit': p.get('unit', ''),
                'price': round(float(price), 2),
                'amount': round(float(amount), 2),
                'dc_unit': dc_unit,
                'dc_total': dc_total,
            })
    return {
        'mode': cur,
        'mode_label': '泰国正报' if cur == 'th' else '香港模式',
        'currency': '美金' if cur == 'th' else '人民币',
        'rate': rate,
        'type_label': _d(meta.get('type_label')),
        'sign_date': _d(meta.get('sign_date')),
        'pickup_date': _d(meta.get('pickup_date')),
        'container_no': _d(meta.get('container_no')),
        'total_amount': round(total_amount, 2),
        'total_declare': round(total_declare, 2),
        'items': items,
    }


def build_all(tax_type_path, out_dir, th_rate=None, buyer=None, mode=None):
    os.makedirs(out_dir, exist_ok=True)
    meta, suppliers = parse_tax_type(tax_type_path)
    # 模式可由网页选择覆盖（香港/泰国），否则按退税类型文件第1行自动判定
    if mode in ('hk', 'th'):
        meta['mode'] = mode
    files = {}
    files['报关单'] = build_customs(meta, suppliers, out_dir, th_rate=th_rate)
    files['申报要素'] = build_declare(suppliers, out_dir)
    files['PI'] = build_pi(meta, suppliers, out_dir, th_rate=th_rate)
    # 采购合同 / 出库单 分开输出（各自复制参考模版并替换标红示例值）
    files['采购合同'] = build_contracts(meta, suppliers, out_dir, buyer=buyer)
    files['出库单'] = build_outbounds(meta, suppliers, out_dir)
    return meta, suppliers, files


if __name__ == '__main__':
    import sys
    src = sys.argv[1] if len(sys.argv) > 1 else r'C:\Users\ASUS\Desktop\退税资料模版\退税类型.xlsx'
    dst = sys.argv[2] if len(sys.argv) > 2 else r'C:\Users\ASUS\WorkBuddy\2026-08-07-14-09-19\退税资料\按退税资料生成\测试输出'
    rate = float(sys.argv[3]) if len(sys.argv) > 3 else None
    m, s, f = build_all(src, dst, th_rate=rate)
    print('模式:', m['mode'], '| 签约:', m['sign_date'], '| 提货:', m['pickup_date'])
    for sup in s:
        print('  供方', sup['name'], '合同号', sup['contract_no'], '产品', len(sup['products']))
    print('生成文件:')
    for k, v in f.items():
        print(' ', k, '->', v)
